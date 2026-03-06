#!/usr/bin/env python3
"""
Convertisseur Rothschild Findings → Word
Génère des documents Word professionnels pour chaque opportunité
"""

import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_rothschild_document(findings_file, output_dir):
    """Crée un document Word à partir des findings de Rothschild"""
    
    # Charger les findings
    with open(findings_file, 'r') as f:
        data = json.load(f)
    
    findings = data.get('findings', [])
    
    if not findings:
        print("❌ Aucun finding trouvé")
        return
    
    for idx, finding in enumerate(findings, 1):
        opp = finding.get('opportunity', {})
        status = finding.get('status', 'unknown')
        confidence = finding.get('confidence', 'medium')
        verdict = finding.get('verdict', '')
        notes = finding.get('notes', '')
        
        opp_id = opp.get('id', f'opportunity-{idx}')
        opp_name = opp.get('name', 'Opportunité sans nom')
        
        # Créer le document
        doc = Document()
        
        # Style par défaut
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # EN-TÊTE
        header = doc.add_heading('ROTHSCHILD RESEARCH DIVISION', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Titre
        title = doc.add_heading(opp_name, 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations clés
        doc.add_heading('📋 INFORMATIONS CLÉS', 2)
        
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('ID Opportunité', opp_id),
            ('Catégorie', opp.get('category', 'N/A')),
            ('Statut', status.upper()),
            ('Confiance', confidence.upper()),
            ('Date validation', finding.get('timestamp', 'N/A')[:10]),
            ('Verdict', verdict)
        ]
        
        for i, (key, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
        
        doc.add_paragraph()
        
        # DESCRIPTION
        doc.add_heading('🎯 DESCRIPTION', 2)
        doc.add_paragraph(opp.get('description', 'Non spécifié'))
        
        # PROBLÈME
        doc.add_heading('❌ PROBLÈME RÉSOLU', 2)
        problem = opp.get('problem', 'Non spécifié')
        doc.add_paragraph(problem)
        
        # SOLUTION
        doc.add_heading('✅ SOLUTION', 2)
        solution = opp.get('solution', 'Non spécifiée')
        doc.add_paragraph(solution)
        
        # MARCHÉ
        doc.add_heading('📊 MARCHÉ', 2)
        
        market_size = opp.get('market_size', {})
        if market_size:
            market_table = doc.add_table(rows=4, cols=2)
            market_table.style = 'Light Grid Accent 1'
            
            market_data = [
                ('TAM (Total Addressable Market)', market_size.get('tam', 'N/A')),
                ('SAM (Serviceable Addressable)', market_size.get('sam', 'N/A')),
                ('SOM (Obtainable Market)', market_size.get('som', 'N/A'))
            ]
            
            for i, (key, value) in enumerate(market_data):
                row = market_table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = value
            
            # Sources
            sources = market_size.get('sources', [])
            if sources:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('Sources: ').bold = True
                p.add_run(', '.join(sources))
        
        # TARGET MARKET
        target_market = opp.get('target_market', {})
        if target_market:
            doc.add_heading('🎯 MARCHÉ CIBLE', 2)
            for key, value in target_market.items():
                p = doc.add_paragraph()
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                p.add_run(value)
        
        # PRICING
        pricing = opp.get('pricing', {})
        if pricing:
            doc.add_heading('💰 PRICING', 2)
            
            price_table = doc.add_table(rows=len(pricing)+1, cols=2)
            price_table.style = 'Light Grid Accent 1'
            
            # Header
            hdr_cells = price_table.rows[0].cells
            hdr_cells[0].text = 'Plan'
            hdr_cells[1].text = 'Prix / Features'
            
            for i, (plan, details) in enumerate(pricing.items(), 1):
                row = price_table.rows[i]
                row.cells[0].text = plan.title()
                row.cells[1].text = str(details)
        
        # UNIT ECONOMICS
        unit_economics = opp.get('unit_economics', {})
        if unit_economics:
            doc.add_heading('📈 UNIT ECONOMICS', 2)
            
            ltv = unit_economics.get('ltv', {})
            cac = unit_economics.get('cac', {})
            
            if ltv:
                p = doc.add_paragraph()
                p.add_run('LTV: ').bold = True
                p.add_run(f"{ltv.get('ltv_raw', 'N/A')} (Ratio {ltv.get('ltv_cac_ratio', 'N/A')}) → Optimisé: {ltv.get('ltv_cac_optimized', 'N/A')}")
            
            if cac:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('CAC par canal:').bold = True
                for channel, cost in cac.items():
                    if channel != 'target_blended':
                        doc.add_paragraph(f"  • {channel}: {cost}", style='List Bullet')
                p = doc.add_paragraph()
                p.add_run(f"CAC Blended Target: {cac.get('target_blended', 'N/A')}").bold = True
        
        # MVP FEATURES
        mvp_features = opp.get('mvp_features', [])
        if mvp_features:
            doc.add_heading('🚀 MVP FEATURES', 2)
            for feature in mvp_features:
                doc.add_paragraph(f"• {feature}", style='List Bullet')
        
        # TECH STACK
        tech_stack = opp.get('tech_stack', {})
        if tech_stack:
            doc.add_heading('🛠️ STACK TECHNIQUE', 2)
            
            tech_table = doc.add_table(rows=len(tech_stack), cols=2)
            tech_table.style = 'Light Grid Accent 1'
            
            for i, (component, details) in enumerate(tech_stack.items()):
                row = tech_table.rows[i]
                row.cells[0].text = component.title()
                row.cells[1].text = str(details)
        
        # COMPETITORS
        competitors = opp.get('competitors', [])
        if competitors:
            doc.add_heading('🏆 CONCURRENTS', 2)
            
            for comp in competitors:
                p = doc.add_paragraph()
                p.add_run(f"{comp.get('name', 'N/A')}").bold = True
                p.add_run(f" - Pricing: {comp.get('pricing', 'N/A')}")
                doc.add_paragraph(f"Faiblesse: {comp.get('weakness', 'N/A')}", style='List Bullet')
        
        # RISKS
        risks = opp.get('risks', [])
        mitigations = opp.get('mitigations', [])
        if risks:
            doc.add_heading('⚠️ RISQUES', 2)
            for i, risk in enumerate(risks):
                doc.add_paragraph(f"{i+1}. {risk}", style='List Number')
            
            if mitigations:
                doc.add_heading('🛡️ MITIGATIONS', 2)
                for i, mitigation in enumerate(mitigations):
                    doc.add_paragraph(f"{i+1}. {mitigation}", style='List Number')
        
        # TIMELINE & COSTS
        doc.add_heading('⏱️ TIMELINE & COÛTS', 2)
        
        timeline_table = doc.add_table(rows=3, cols=2)
        timeline_table.style = 'Light Grid Accent 1'
        
        timeline_data = [
            ('Coût de démarrage', opp.get('startup_cost', 'N/A')),
            ('Time to Launch', opp.get('time_to_launch', 'N/A'))
        ]
        
        for i, (key, value) in enumerate(timeline_data):
            row = timeline_table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
        
        # REVENUE POTENTIAL
        revenue = opp.get('revenue_potential', {})
        if revenue:
            doc.add_heading('💵 POTENTIEL DE REVENUS', 2)
            for scenario, value in revenue.items():
                p = doc.add_paragraph()
                p.add_run(f"{scenario.title()}: ").bold = True
                p.add_run(value)
        
        # NOTES ROTHCHILD
        if notes:
            doc.add_heading('📝 NOTES ROTHCHILD', 2)
            note_para = doc.add_paragraph(notes)
            note_para.italic = True
        
        # FOOTER
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f"\nDocument généré par Rothschild Research Division — {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sauvegarder
        filename = f"{opp_id}-rothschild-findings.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        
        print(f"✅ Document créé: {filename}")
    
    print(f"\n🎉 {len(findings)} document(s) Word généré(s) dans: {output_dir}")

if __name__ == '__main__':
    findings_file = '/root/.openclaw/workspace/mission-control/data/rothschild-raw-findings.json'
    output_dir = '/root/.openclaw/workspace/mission-control/business-plans/word'
    
    os.makedirs(output_dir, exist_ok=True)
    create_rothschild_document(findings_file, output_dir)
